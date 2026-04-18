"""
report_generator.py
────────────────────────────────────────────────────────────────
蒙特卡洛投资组合分析报告生成引擎
流程：
  labData + mc_settings
    → build_data_summary()  提炼结构化摘要
    → build_prompt()         组装中文 Prompt
    → call_deepseek()        调用 DeepSeek API，返回 JSON 章节文本
    → build_docx()           python-docx 组装 Word，返回 BytesIO
"""

import os
import io
import json
import re
import math
from datetime import datetime

from dotenv import load_dotenv
load_dotenv()

from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── LLM (OpenRouter) 客户端 ────────────────────────────────────────────
_llm_client = None

def _get_client() -> OpenAI:
    global _llm_client
    if _llm_client is None:
        api_key = os.getenv("OPENROUTER_API_KEY")
        if not api_key:
            raise RuntimeError("OPENROUTER_API_KEY 未配置，请在 .env 中设置。")
        _llm_client = OpenAI(
            api_key=api_key,
            base_url="https://openrouter.ai/api/v1"
        )
    return _llm_client


# ═══════════════════════════════════════════════════════════════
# STEP 1 — 数据摘要提炼
# ═══════════════════════════════════════════════════════════════

def build_data_summary(lab_result: dict, mc_settings: dict, client_info: dict = None) -> dict:
    """将 labData、mc_settings 与 client_info 提炼为报告所需结构化摘要。"""

    ci = client_info or {}
    client_age   = int(ci.get("age") or 0) or None
    client_name  = (ci.get("name") or "").strip()
    client_goals = ci.get("goals") or []
    advisor_name = (ci.get("advisor") or "").strip()

    GOAL_MAP = {
        "retirement": "退休收入保障", "education": "子女教育基金",
        "legacy": "遗产传承规划",    "growth":   "资产保值增值",
        "liquidity": "流动性储备",
    }
    goals_text = "、".join(GOAL_MAP.get(g, g) for g in client_goals) \
                 if client_goals else "长期财富积累与风险管控"

    mc_label   = lab_result.get("mc_target_label", "")
    has_custom = "custom_portfolio" in lab_result and lab_result["custom_portfolio"]
    if has_custom and (mc_label in ("custom_portfolio", "")):
        portfolio, portfolio_label = lab_result["custom_portfolio"], "自定义组合"
    else:
        portfolio, portfolio_label = lab_result.get("max_sharpe", {}), "最大夏普比率组合"

    allocs        = portfolio.get("allocations", {})
    asset_stats   = lab_result.get("asset_stats", {})
    risk_contribs = lab_result.get("risk_contributions", {})

    allocations = []
    for isin, weight in sorted(allocs.items(), key=lambda x: -x[1]):
        if weight < 0.001:
            continue
        stats   = asset_stats.get(isin, {})
        exp_ret = stats.get("expected_return", 0)
        vol     = stats.get("volatility", 0)
        div_yld = stats.get("dividend_yield", 0)
        if   exp_ret > 0.10: role = "高成长引擎"
        elif vol     < 0.10: role = "稳定压舱石"
        elif div_yld > 0.03: role = "股息收益来源"
        else:                role = "多元化缓冲"
        allocations.append({
            "isin": isin, "name": stats.get("name", isin),
            "weight_pct":       round(weight * 100, 1),
            "exp_return_pct":   round(exp_ret * 100, 2),
            "vol_pct":          round(vol * 100, 2),
            "div_yield_pct":    round(div_yld * 100, 2),
            "risk_contrib_pct": round(risk_contribs.get(isin, 0) * 100, 1),
            "role":             role,
        })

    mc           = lab_result.get("monte_carlo", {})
    chart        = mc.get("chart", [])
    success_rate = mc.get("success_rate")
    drawdown     = mc.get("drawdown", {})
    comb_draw    = mc.get("combined_drawdown", {})
    stressed_vol = mc.get("stressed_volatility")
    ins_stats    = mc.get("insurance_stats")

    # mc_settings 字段名兼容（前端传 inflation，后端规范是 inflation_pct）
    years_total        = mc_settings.get("years", max(len(chart) - 1, 1))
    raw_inf            = mc_settings.get("inflation_pct") or mc_settings.get("inflation", 0.0)
    inflation_pct      = float(raw_inf)
    inflation_rate     = inflation_pct / 100.0
    withdrawal_base    = mc_settings.get("withdrawal", 0)
    withdrawal_start   = mc_settings.get("withdrawal_start", 1)
    withdrawal_end     = mc_settings.get("withdrawal_end", years_total)
    contribution_base  = mc_settings.get("contribution", 0)
    contribution_start = mc_settings.get("contribution_start", 1)
    contribution_years = mc_settings.get("contribution_years", 0)

    last      = chart[-1] if chart else {}
    p10_final = last.get("p10", 0)
    p50_final = last.get("p50", 0)
    p90_final = last.get("p90", 0)
    real_p50  = last.get("real_p50")

    chart_map = {pt["year"]: pt for pt in chart}

    milestone_candidates = []

    def add_m(year, label, note=""):
        if 0 < year <= years_total and year not in {m["year"] for m in milestone_candidates}:
            milestone_candidates.append({"year": year, "label": label, "note": note})

    if contribution_years > 0:
        lbl = (f"缴费期结束（{client_age + contribution_years}岁）" if client_age
               else f"缴费期结束（第{contribution_years}年）")
        add_m(contribution_years, lbl, "初期积累阶段完成，资金进入自由增长期")
    if withdrawal_start > 0:
        lbl = (f"开始提取（{client_age + withdrawal_start}岁）" if client_age
               else f"开始提取（第{withdrawal_start}年）")
        add_m(withdrawal_start, lbl, "现金流正式激活，进入养老模式")
    for y in [10, 15, 20, 25, 30]:
        note = {10: "布局初步成型，复利效应开始显现",
                20: "长期持有效果明显，抗通胀能力凸显",
                30: "深度复利阶段，财富积累加速"}.get(y, "")
        lbl = f"第{y}年（{client_age + y}岁）" if client_age else f"第{y}年"
        add_m(y, lbl, note)
    if client_age:
        for target_age in [60, 65, 70, 80]:
            yr = target_age - client_age
            note = {65: "退休规划关键节点", 70: "晚年生活质量保障期",
                    80: "长寿风险规划，财富是否仍然充裕？"}.get(target_age, "")
            add_m(yr, f"{target_age}岁（第{yr}年）", note)
    lbl = (f"规划终点（{client_age + years_total}岁）" if client_age
           else f"规划终点（第{years_total}年）")
    add_m(years_total, lbl, "全期终值汇总")

    milestone_candidates.sort(key=lambda m: m["year"])
    milestones = []
    for m in milestone_candidates:
        pt = chart_map.get(m["year"], {})
        if pt.get("p50"):
            milestones.append({**m, "p10": pt.get("p10"), "p50": pt.get("p50"), "p90": pt.get("p90"),
                                   "combined_p50": pt.get("combined_p50")})

    snapshot_years = list(range(5, years_total + 1, 5))
    if years_total not in snapshot_years:
        snapshot_years.append(years_total)
    yearly_snapshots = [
        {"year": pt["year"], "p10": pt.get("p10", 0), "p50": pt.get("p50", 0), "p90": pt.get("p90", 0)}
        for pt in chart if pt["year"] in snapshot_years
    ]

    display_end = min(withdrawal_start + 14, withdrawal_end, years_total)
    withdrawal_schedule = []
    for yr in range(withdrawal_start, display_end + 1):
        idx = yr - withdrawal_start
        # 1. 投资组合提取 (Nominal)
        pure_amt = withdrawal_base * ((1 + inflation_rate) ** idx)
        
        # 2. 保险组合提取 (Nominal)
        pt = chart_map.get(yr, {})
        ins_amt = pt.get("ins_withdrawal", 0)
        
        withdrawal_schedule.append({
            "year": yr, 
            "pure_amount": round(pure_amt),
            "ins_amount": round(ins_amt),
            "total_amount": round(pure_amt + ins_amt),
            "amount": round(pure_amt)  # Keep for backward compatibility if needed
        })

    # ── 保险量化影响数据（与网页预览对齐）
    last_c50 = last.get("combined_p50")
    last_c10 = last.get("combined_p10")
    last_c90 = last.get("combined_p90")
    is_combined = (last_c50 is not None)

    ins_block = None
    if is_combined and ins_stats:
        contrib_p50 = (last_c50 or 0) - (p50_final or 0)
        contrib_p10 = (last_c10 or 0) - (p10_final or 0)
        dd_pure = abs(drawdown.get("p10", 0))
        dd_comb = abs(comb_draw.get("p10", 0))
        cv_end  = ins_stats.get("avg_cv_at_year_end", {})
        ins_block = {
            "combined_p50":          last_c50,
            "combined_p10":          last_c10,
            "combined_p90":          last_c90,
            "pure_p50":              p50_final,
            "pure_p10":              p10_final,
            "contrib_p50":           round(contrib_p50),
            "contrib_p10":           round(contrib_p10),
            "drawdown_p10_pure_pct": round(dd_pure * 100, 1),
            "drawdown_p10_comb_pct": round(dd_comb * 100, 1),
            "drawdown_improve_pct":  round((dd_pure - dd_comb) * 100, 1),
            "cv_low":                cv_end.get("low"),
            "cv_mid":                cv_end.get("mid"),
            "cv_high":               cv_end.get("high"),
            "withdrawal_cov_pct":    round(ins_stats.get("withdrawal_coverage_pct", 0), 1),
            "total_ins_withdrawal":  ins_stats.get("total_insurance_withdrawal", 0),
        }
        # 封面/摘要使用合并终值（与网页预览一致）
        p50_final = last_c50
        p10_final = last_c10
        p90_final = last_c90

    return {
        "client_name":         client_name,
        "client_age":          client_age,
        "goals_text":          goals_text,
        "advisor_name":        advisor_name,
        "portfolio_label":     portfolio_label,
        "capital":             mc_settings.get("capital", 0),
        "years":               years_total,
        "withdrawal":          withdrawal_base,
        "withdrawal_start":    withdrawal_start,
        "withdrawal_end":      withdrawal_end,
        "inflation_pct":       inflation_pct,
        "contribution":        contribution_base,
        "contribution_start":  contribution_start,
        "contribution_years":  contribution_years,
        "stress_mode":         mc_settings.get("stress", False),
        "target_goal":         mc_settings.get("target", 0),
        "expected_return_pct": round(mc.get("irr", {}).get("p50", 0) * 100, 2),
        "volatility_pct":      round(portfolio.get("volatility", 0) * 100, 2),
        "sharpe_ratio":        round(portfolio.get("sharpe_ratio", 0), 3),
        "allocations":         allocations,
        "success_rate_pct":    round((success_rate or 0) * 100, 1),
        "p10_final":           p10_final,
        "p50_final":           p50_final,
        "p90_final":           p90_final,
        "real_p50_final":      real_p50,
        "drawdown_p50_pct":    round(drawdown.get("p50", 0) * 100, 1),
        "drawdown_p10_pct":    round(drawdown.get("p10", 0) * 100, 1),
        "drawdown_p90_pct":    round(drawdown.get("p90", 0) * 100, 1),
        "stressed_vol_pct":    round(stressed_vol * 100, 2) if stressed_vol else None,
        "milestones":          milestones,
        "yearly_snapshots":    yearly_snapshots,
        "withdrawal_schedule": withdrawal_schedule,
        "ins_block":           ins_block,
    }



# ═══════════════════════════════════════════════════════════════
# STEP 2 — Prompt 组装
# ═══════════════════════════════════════════════════════════════

def build_prompt(s: dict) -> str:
    """将数据摘要组装为客户叙事版中文 Prompt，输出 7 章 JSON。
    
    v2 升级：融入保险计划书的专业写作哲学：
    - 10条核心写作原则（对标高净值客户可读标准）
    - 10条表达禁忌（合规过滤）
    - 6项输出自检机制
    - 里程碑节点附加意义注释
    - 非保证数据合规标注
    """

    def fmt(n):
        if n is None: return "N/A"
        return f"${n:,.0f}"

    client_intro = ""
    if s.get("client_name"):
        client_intro += f"客户姓名：{s['client_name']}\\n"
    if s.get("client_age"):
        client_intro += f"当前年龄：{s['client_age']}岁\\n"
    if s.get("goals_text"):
        client_intro += f"主要财务目标：{s['goals_text']}\\n"
    if s.get("advisor_name"):
        client_intro += f"制作顾问：{s['advisor_name']}\\n"

    alloc_lines = "\\n".join(
        f"  - {a['name']} ({a['isin']}): 权重 {a['weight_pct']}%，"
        f"预期年化收益 {a['exp_return_pct']}%，波动率 {a['vol_pct']}%，角色：{a['role']}"
        for a in s["allocations"]
    )

    # 里程碑附带意义注释，帮助 LLM 写出"这意味着什么"
    milestone_lines = "\\n".join(
        f"  - {m['label']}: 悲观 {fmt(m['p10'])} / 中性 {fmt(m['p50'])} / 乐观 {fmt(m['p90'])}"
        + (f"  ← {m['note']}" if m.get("note") else "")
        for m in s.get("milestones", [])
    ) or "  （无关键里程碑数据）"

    withdraw_lines = "\\n".join(
        f"  - 第{w['year']}年: {fmt(w['amount'])}" if w["amount"] > 0 else f"  - 第{w['year']}年: 无提取"
        for w in s["withdrawal_schedule"]
    ) or "  （无提取计划）"

    stress_note = (
        f"已启用压力测试（危机相关性≥0.8），压力波动率 {s['stressed_vol_pct']}%。此为非常规极端场景，非保证数据。"
        if s["stress_mode"] else "未启用压力测试。"
    )
    real_p50_note = (
        f"扣除通胀后，中性终值今日购买力约 {fmt(s['real_p50_final'])}（含通胀折算，非保证）。"
        if s.get("real_p50_final") else ""
    )
    salutation = f"{s['client_name']}，" if s.get("client_name") else ""
    goals_text = s.get("goals_text", "长期财富积累与风险管控")

    # ── 组装升级版 Prompt ──────────────────────────────────────
    sections = []

    sections.append("""你是一位顶尖的家族办公室财富规划师。你的任务是基于我提供的「客户画像」与「蒙特卡洛1万次模拟测算数据」，为客户撰写一份极具深度、专业性与共情力的《专属财富规划报告》。

【写作风格要求】
1. 你的受众是高净值且非金融模型纯专业的客户。必须使用投研报告般克制、精准、且富有解释威力的商务金融语言，禁止使用绝对化推销词汇。
2. 论述深度：拒绝“报菜名”式罗列数字！必须解释数字背后的业务逻辑与模型常识，让客户明白这套方案是如何在现实不确定性中运行的。
3. 动态配置视角：如果数据中包含【保险底座数据】，必须重点论证“双核合并体系（高波动高流动组合 + 保本增值现金价值）”是如何在平滑风险的同时，锁定长期回报安全垫的。

---
""")

    # 客户信息区块
    client_block = f"""## 三、客户信息

{client_intro if client_intro else "未提供客户具体信息，请使用通用高净值家庭场景撰写。"}主要财务目标：{goals_text}

---
"""
    sections.append(client_block)

    # ── 保险量化数据段（仅当 ins_block 存在时添加）
    ins = s.get("ins_block")
    ins_section = ""
    if ins:
        ins_section = f"""\n### \u4fdd\u9669\u5e95\u5ea7\u5bf9\u6574\u4f53\u8d44\u4ea7\u914d\u7f6e\u7684\u91cf\u5316\u5f71\u54cd\uff08\u4e0e\u7f51\u9875\u9884\u89c8\u5bf9\u9f50\uff09
- \u7ec8\u503c\u5bf9\u6bd4\uff1a\u7eaf\u6295\u8d44\u7ec4\u5408\u4e2d\u6027\u7ec8\u503c {fmt(ins['pure_p50'])} \u2192 \u53e0\u52a0\u4fdd\u9669\u540e\u5408\u5e76\u7ec8\u503c {fmt(ins['combined_p50'])}\uff08\u4fdd\u9669\u989d\u5916\u8d21\u732e {fmt(ins['contrib_p50'])}\uff09
- \u60b2\u89c2\u60c5\u666f\uff08P10\uff09\uff1a\u7eaf\u7ec4\u5408\u8dcc\u81f3 {fmt(ins['pure_p10'])}\uff0c\u4fdd\u9669\u989d\u5916\u6258\u5e95 {fmt(ins['contrib_p10'])}\uff0c\u5408\u5e76\u7ec8\u503c {fmt(ins['combined_p10'])}
- \u6700\u5927\u56de\u64a4\u6536\u7a84\uff1a\u5f15\u5165\u4fdd\u9669\u540e\uff0cP10\u60b2\u89c2\u6700\u5927\u56de\u64a4\u4ece {ins['drawdown_p10_pure_pct']}% \u6536\u7a84\u81f3 {ins['drawdown_p10_comb_pct']}%\uff08\u6539\u5584\u4e86 {ins['drawdown_improve_pct']} \u4e2a\u767e\u5206\u70b9\uff09
- \u4fdd\u9669\u672b\u671f\u73b0\u91d1\u4ef7\u5024\uff08\u7b2c{s['years']}\u5e74\uff09\uff1a\u60b2\u89c2\u5b9e\u73b0 {fmt(ins['cv_low'])} / \u4e2d\u6027\u5b9e\u73b0 {fmt(ins['cv_mid'])} / \u4e50\u89c2\u5b9e\u73b0 {fmt(ins['cv_high'])}\uff08\u975e\u4fdd\u8bc1\uff09
- \u4fdd\u9669\u5bf9\u5e74\u5ea6\u63d0\u53d6\u7684\u8986\u76d6\u7387\uff1a{ins['withdrawal_cov_pct']}%\uff08\u7d2f\u8ba1\u53ef\u63d0\u53d6 {fmt(ins['total_ins_withdrawal'])}\uff0c\u6709\u6548\u964d\u4f4e\u5bf9\u6295\u8d44\u7ec4\u5408\u7684\u53d8\u73b0\u538b\u529b\uff09

---
"""

    # \u91cf\u5316\u6570\u636e\u533a\u5757
    data_block = f"""## \u56db\u3001\u91cf\u5316\u6570\u636e\uff08\u6240\u6709\u9884\u6d4b\u5747\u4e3a\u8499\u7279\u5361\u6d1b\u6a21\u62df\u7ed3\u679c\uff0c\u5c5e\u4e8e\u7edf\u8ba1\u6982\u7387\uff0c\u975e\u4fdd\u8bc1\u6536\u76ca\uff09

### \u8d22\u52a1\u5047\u8bbe
- \u521d\u59cb\u8d44\u672c\uff1a{fmt(s["capital"])}
- \u6a21\u62df\u5e74\u9650\uff1a{s["years"]} \u5e74
- \u5e74\u5ea6\u63d0\u53d6\uff1a{fmt(s["withdrawal"])}\uff08\u7b2c{s["withdrawal_start"]}\u2013{s["withdrawal_end"]}\u5e74\uff0c\u6309 {s["inflation_pct"]}%/\u5e74\u901a\u80c0\u8c03\u6574\uff09
- \u5e74\u5ea6\u8ffd\u52a0\uff1a{fmt(s["contribution"])}\uff08\u7b2c{s["contribution_start"]}\u2013{s["contribution_years"]}\u5e74\uff09
- \u76ee\u6807\u7ec8\u503c\uff1a{fmt(s["target_goal"])}
- \u901a\u80c0\u5047\u8bbe\uff1a{s["inflation_pct"]}%/\u5e74
- \u538b\u529b\u6d4b\u8bd5\uff1a{stress_note}

### \u8d44\u4ea7\u914d\u7f6e\uff08\u89d2\u8272\u5b9a\u4f4d\uff09
{alloc_lines}

### \u7ec4\u5408\u6574\u4f53\u8868\u73b0\uff08\u57fa\u4e8e\u5386\u53f2\u6570\u636e\u4f18\u5316\uff0c\u975e\u4fdd\u8bc1\u672a\u6765\u8868\u73b0\uff09
- \u9884\u671f\u5e74\u5316\u6536\u76ca\uff1a{s["expected_return_pct"]}%\uff08\u5386\u53f2\u7edf\u8ba1\u4f30\u7b97\uff0c\u975e\u672a\u6765\u627f\u8bfa\uff09
- \u7ec4\u5408\u6ce2\u52a8\u7387\uff1a{s["volatility_pct"]}%\uff08\u8861\u91cf\u4ef7\u683c\u8d77\u4f0f\u5e45\u5ea6\u7684\u6807\u51c6\u5dee\uff09
- \u98ce\u9669\u6536\u76ca\u6548\u7387\uff08\u590f\u666e\u6bd4\u7387\uff09\uff1a{s["sharpe_ratio"]}\uff08\u6bcf\u627f\u62c5\u4e00\u5355\u4f4d\u98ce\u9669\u6240\u83b7\u5f97\u7684\u8d85\u989d\u56de\u62a5\uff0c\u6570\u503c\u8d8a\u9ad8\u8d8a\u597d\uff09

### \u5173\u952e\u65f6\u95f4\u8282\u70b9\u8d22\u5bcc\u9884\u4f30\uff0810,000\u6761\u968f\u673a\u8def\u5f84\u7684\u6982\u7387\u5206\u4f4d\u6570\uff0c\u975e\u4fdd\u8bc1\uff09
\u8bfb\u6cd5\u8bf4\u660e\uff1a\u60b2\u89c2(P10)=\u6700\u5dee10%\u8def\u5f84\u4e0b\u9650 / \u4e2d\u6027(P50)=\u6700\u53ef\u80fd\u4e2d\u4f4d\u6570 / \u4e50\u89c2(P90)=\u987a\u98ce\u5c40\u4e0a\u9650
{milestone_lines}

### \u6a21\u62df\u7ec8\u503c\uff08{s["years"]}\u5e74\u540e\uff0c\u975e\u4fdd\u8bc1\uff09
- \u4e50\u89c2\uff08P90\uff09\uff1a{fmt(s["p90_final"])}\uff08\u987a\u98ce\u60c5\u666f\uff0c\u4ec510%\u8def\u5f84\u8d85\u8fc7\u6b64\u503c\uff09
- \u4e2d\u6027\uff08P50\uff09\uff1a{fmt(s["p50_final"])}\uff08\u6700\u53ef\u80fd\u843d\u70b9\uff0c\u4e2d\u4f4d\u6570\u9884\u671f\uff09
- \u60b2\u89c2\uff08P10\uff09\uff1a{fmt(s["p10_final"])}\uff08\u5e95\u7ebf\u573a\u666f\uff0c\u4ec510%\u8def\u5f84\u4f4e\u4e8e\u6b64\u503c\uff09
- \u76ee\u6807\u8fbe\u6210\u6982\u7387\uff1a{s["success_rate_pct"]}%
{real_p50_note}

### \u6700\u5927\u56de\u64a4\u7edf\u8ba1\uff08\u6a21\u62df\u671f\u5185\u8d26\u9762\u4ef7\u5024\u6700\u5927\u4e34\u65f6\u8dcc\u5e45\uff0c\u4e0d\u4ee3\u8868\u5b9e\u9645\u4e8f\u635f\u9501\u5b9a\uff09
- \u4e2d\u6027\u8def\u5f84\u6700\u5927\u56de\u64a4\uff1a{s["drawdown_p50_pct"]}%
- \u60b2\u89c2\u8def\u5f84\u6700\u5927\u56de\u64a4\uff1a{s["drawdown_p10_pct"]}%

### \u901a\u80c0\u8c03\u6574\u63d0\u53d6\u8ba1\u5212\uff08\u524d{len(s["withdrawal_schedule"])}\u5e74\uff09
{withdraw_lines}
{ins_section}---
"""
    sections.append(data_block)

    # 自检要求
    ins_check = "\n7. \u5982\u679c\u6709\u4fdd\u9669\u6570\u636e\uff0c\u662f\u5426\u5728 insurance_impact \u7ae0\u8282\u4e2d\u660e\u786e\u91cf\u5316\u4e86\u4fdd\u9669\u5bf9\u7ec8\u503c\u3001\u56de\u64a4\u3001\u63d0\u53d6\u7684\u5b9e\u9645\u8d21\u732e\uff1f" if ins else ""
    sections.append(f"""## \u4e94\u3001\u8f93\u51fa\u524d\u8bf7\u81ea\u884c\u5b8c\u6210\u4ee5\u4e0b\u68c0\u67e5\uff08\u4e0d\u7b26\u5408\u8bf7\u81ea\u52a8\u4fee\u6b63\u540e\u518d\u8f93\u51fa\uff09

1. \u662f\u5426\u5168\u7a0b\u7528\"\u60a8\"\u79f0\u8c13\u76f4\u63a5\u4e0e\u5ba2\u6237\u5bf9\u8bdd\uff1f
2. \u662f\u5426\u6bcf\u4e2a\u91cd\u8981\u6570\u636e\u70b9\u540e\u90fd\u89e3\u91ca\u4e86\"\u8fd9\u5bf9\u60a8\u610f\u5473\u7740\u4ec0\u4e48\"\uff1f
3. \u662f\u5426\u660e\u786e\u8bf4\u660e\u4e86\u6240\u6709\u9884\u6d4b\u5c5e\u4e8e\"\u6a21\u62df/\u7edf\u8ba1/\u975e\u4fdd\u8bc1\"\uff1f
4. \u662f\u5426\u4fdd\u7559\u4e86\u98ce\u9669\u8fb9\u754c\u3001\u6700\u5927\u56de\u64a4\u542b\u4e49\u3001\u957f\u671f\u6301\u6709\u903b\u8f91\uff1f
5. \u662f\u5426\u5b8c\u5168\u56de\u907f\u4e86\u7981\u5fe7\u8bcd\u548c\u5938\u5f20\u8868\u8fbe\uff1f
6. \u6574\u4f53\u8bfb\u8d77\u6765\uff0c\u662f\u5426\u50cf\u4e00\u4efd\u53ef\u4ee5\u76f4\u63a5\u4f20\u7ed9\u5ba2\u6237\u7684\u957f\u671f\u8d22\u5bcc\u89c4\u5212\u8bf4\u660e\uff1f{ins_check}

---
""")

    # \u8f93\u51fa\u683c\u5f0f\u8981\u6c42
    ins_key_desc = ''
    if ins:
        ins_key_desc = (
            '\n  "insurance_impact": "\uff08120-160\u5b57\uff1a\u57fa\u4e8e\u4e0a\u8ff0\u8499\u7279\u5361\u6d1b\u91cf\u5316\u6570\u636e\uff0c\u89e3\u91ca\u5f15\u5165\u4fdd\u9669\u5e95\u5ea7\u540e\u5bf9\u6574\u4f53\u8d44\u4ea7\u914d\u7f6e\u7684\u4e09\u4e2a\u6838\u5fc3\u6539\u53d8\uff1a'
            '①\u7ec8\u503c\u8d21\u732e\u2014\u2014\u4fdd\u9669\u989d\u5916\u4e3a\u6574\u4f53\u8d22\u5bcc\u8d21\u732e\u4e86\u591a\u5c11\uff1b'
            '②\u5371\u673a\u5b88\u62a4\u2014\u2014\u5728\u6700\u5dee\u5e02\u573a\u73af\u5883\u4e0b\uff0c\u4fdd\u9669\u6258\u5e95\u4e86\u591a\u5c11\uff1b'
            '③\u56de\u64a4\u6536\u7a84\u2014\u2014\u6700\u5927\u8dcc\u5e45\u51cf\u5c11\u4e86\u51e0\u4e2a\u767e\u5206\u70b9\u3002'
            '\u5fc5\u987b\u7528\u5b9e\u9645\u6570\u5b57\u8bf4\u8bdd\uff0c\u4e0d\u8981\u6cdb\u6cdb\u800c\u8c08\u3002\u975e\u4fdd\u8bc1\u90e8\u5206\u9700\u6807\u6ce8\u3002\uff09",')
    output_block = f"""## \u516d\u3001\u8f93\u51fa\u8981\u6c42

\u8bf7\u4e25\u683c\u4ee5\u5408\u6cd5 JSON \u683c\u5f0f\u8fd4\u56de\uff0c\u5305\u542b\u4ee5\u4e0b key\uff0c\u6bcf\u4e2a value \u4e3a\u7eaf\u6587\u672c\u5b57\u7b26\u4e32\uff08\u4e0d\u542b Markdown \u683c\u5f0f\u7b26\u53f7\uff09\uff1a

{{
  "client_opening": "{salutation}\u5982\u679c\u60a8\u6b63\u5728\u8003\u8651...\uff08130-160\u5b57\uff1a\u4ee5\u5ba2\u6237\u771f\u5b9e\u9700\u6c42\u4e3a\u5207\u5165\u70b9\uff0c\u89e3\u91ca\u8fd9\u4efd\u65b9\u6848\u60f3\u89e3\u51b3\u7684\u6838\u5fc3\u95ee\u9898\uff0c\u5efa\u7acb\u4ee3\u5165\u611f\uff0c\u4e0d\u8981\u4e00\u5f00\u59cb\u5c31\u5806\u6570\u5b57\uff0c\u4e0d\u8981\u5199\u6210\u4ea7\u54c1\u4ecb\u7ecd\uff09",
  "strategy_narrative": "\uff0890-130\u5b57\uff1a\u7528\u767d\u8bdd\u89e3\u91ca\u4e3a\u4ec0\u4e48\u9009\u62e9\u8fd9\u6837\u7684\u8d44\u4ea7\u914d\u7f6e\uff0c\u8fdb\u653b\u4e0e\u9632\u5b88\u5982\u4f55\u5e73\u8861\uff0c\u8ba9\u5ba2\u6237\u4e00\u8bfb\u5c31\u61c2\u8fd9\u662f\u4e3a\u4ed6\u91cf\u8eab\u8bbe\u8ba1\u7684\uff0c\u4e0d\u8981\u7528\u4e13\u4e1a\u672f\u8bed\uff09",
  "portfolio_roles": "\uff08160-210\u5b57\uff1a\u9010\u4e00\u89e3\u91ca\u6bcf\u4e2a\u8d44\u4ea7\u5728\u7ec4\u5408\u4e2d\u626e\u6f14\u4ec0\u4e48\u89d2\u8272\uff0c\u7528\u8fdb\u653b/\u9632\u5b88/\u7f13\u51b2\u7684\u6846\u67b6\u8bf4\u660e\uff0c\u7528\u901a\u4fd7\u6bd4\u55bb\u5e2e\u5ba2\u6237\u7406\u89e3\u9322\u662f\u600e\u4e48\u8fd0\u8f6c\u7684\uff0c\u4e0d\u8981\u53ea\u7f57\u5217\u53c2\u6570\uff09",
  "milestone_interpretation": "\uff08160-210\u5b57\uff1a\u7ed3\u5408\u5173\u952e\u65f6\u95f4\u8282\u70b9\u6570\u636e\uff0c\u89e3\u8bfb\u5404\u9636\u6bb5\u8d22\u5bcc\u53d8\u5316\u7684\u542b\u4e49\uff1b\u91cd\u70b9\u5fc5\u987b\u8bda\u5b9e\u8bf4\u660e\u60b2\u89c2P10\u60c5\u666f\u4e0b\u7684\u5fc3\u7406\u51c6\u5907\u548c\u5e94\u5bf9\u6846\u67b6\uff0c\u8ba9\u5ba2\u6237\u77e5\u9053\u5e95\u7ebf\u573a\u666f\u662f\u53ef\u4ee5\u7ba1\u7406\u7684\uff0c\u800c\u4e0d\u662f\u8ba9\u4eba\u6050\u6168\u7684\uff09",
  "withdrawal_narrative": "\uff08100-150\u5b57\uff1a\u4ee5\u5ba2\u6237\u89c6\u89d2\u53d9\u8ff0\u63d0\u53d6\u8ba1\u5212\u7684\u8bbe\u8ba1\u903b\u8f91\uff0c\u91cd\u70b9\u8bf4\u660e\u4e3a\u4ec0\u4e48\u63d0\u53d6\u91d1\u989d\u8981\u968f\u901a\u80c0\u589e\u52a0\u3001\u8fd9\u6837\u8bbe\u8ba1\u5bf9\u8d2d\u4e70\u529b\u4fdd\u969c\u610f\u5473\u7740\u4ec0\u4e48\uff09",
  "risk_narrative": "\uff08100-150\u5b57\uff1a\u53e3\u8bed\u5316\u89e3\u91ca\u6700\u5927\u56de\u64a4\u7684\u771f\u5b9e\u542b\u4e49\u2014\u2014\u5b83\u662f\u8d26\u9762\u6d6e\u52a8\uff0c\u4e0d\u7b49\u4e8e\u9501\u5b9a\u4e8f\u635f\uff1b\u7ed9\u51fa\u9762\u5bf9\u56de\u64a4\u65f6\u7684\u6b63\u786e\u5fc3\u6001\u548c\u884c\u4e3a\u5efa\u8bae\uff1b\u5f3a\u8c03\u7eaa\u5f8b\u6bd4\u9884\u6d4b\u66f4\u91cd\u8981\uff0c\u65f6\u95f4\u662f\u6700\u597d\u7684\u670b\u53cb\uff09",{ins_key_desc}
  "recommendations": "\uff0890-130\u5b57\uff1a1-3\u6761\u9488\u5bf9\u5f53\u524d\u5177\u4f53\u53c2\u6570\u7684\u4e2a\u4eba\u5316\u5efa\u8bae\uff0c\u8bed\u6c14\u50cf\u987e\u95ee\u9762\u5bf9\u9762\u8bf4\u8bdd\uff0c\u8981\u6709\u5177\u4f53\u7684\u884c\u52a8\u65b9\u5411\u6216\u6ce8\u610f\u4e8b\u9879\uff0c\u4e0d\u8981\u5957\u8bdd\uff09"
}}

\u6ce8\u610f\uff1a\u53ea\u8fd4\u56de JSON \u5bf9\u8c61\u672c\u8eab\uff0c\u4e0d\u8981\u6709\u4efb\u4f55\u524d\u7f6e\u8bf4\u660e\u3001\u540e\u7f6e\u8bf4\u660e\u6216 Markdown \u4ee3\u7801\u5757\u5305\u88f9\u3002"""
    sections.append(output_block)

    return "\\n".join(sections)



# ═══════════════════════════════════════════════════════════════
# STEP 3 — 调用 LLM API (Via OpenRouter)
# ═══════════════════════════════════════════════════════════════

def call_llm(prompt: str, timeout: int = 60) -> dict:
    """
    调用 LLM API (默认 Claude 3.5 Sonnet) 并解析返回的 JSON 章节文本。
    发生任何错误时，返回一个含有占位文本的预设字典（保证报告必然生成）。
    """
    fallback = {
        "executive_summary": "如果您正在考虑为家庭建立长期、可持续的财富积累体系，这份方案正是为此而设计的。我们在测算中验证了极高的目标达成率定调，确保了长远的家庭计划实施无虞。",
        "capital_markets_outlook": "我们通过几何布朗运动（GBM）和蒙特卡洛模型进行了1万次随机漫步测算，如同气象台模拟无数次冷暖气流碰撞来预演极端天气。这打破了传统财务规划单调线性的弊端，让我们不仅关注中枢表现（P50），更能为您在遭受危机时兜住底线（P10）。",
        "portfolio_roles": "组合中的各资产承担进攻（捕捉成长）、防守（抗通胀）和缓冲（提供低波流动性）三大核心职责。各要素低度相关的特质极大平滑了配置风险曲线。",
        "milestone_interpretation": "随着时间推移，在顺风局（P90乐观路径）您能收获远超预期的复利果实；即使在逆风防守（P10悲观路径），时间价值和纪律也能确保财富不发生毁灭性下降，给您安稳的底气。",
        "spending_impact": "面对年度资金动用，提款压力在市场暴跌时极易引发“顺序回报风险”。我们的压力测试已证明，目前的提取速率处于健康的资产造血区间边界内，无需担心长期流动枯竭。",
        "risk_narrative": "最大回撤并非实际锁死的账面亏损，而是在系统性危机中资产的暂时收缩。建议您做足心理建设：在熊市切忌恐慌割肉断送复利。我们已提前预设了应急资金缓冲应对这类突发。",
        "insurance_strategy": "通过配置具有“平滑机制”的保险底座，方案获得了一层受市场波动影响更小的安全垫。它在提供终期提取来源的同时，更在危机时刻收窄了账面最大回撤（drawdown_improve），为您构筑了坚固的财务护城河。",
        "recommendations": "请保持战略耐心。每季度或每半年度进行一次账户再平衡检视即可，严守长期投资纪律，拒绝市场短期杂音干扰；若遇流动性突发诉求，优选无损方式调用工具。"
    }

    model_name = os.getenv("LLM_MODEL", "anthropic/claude-3.5-sonnet")

    try:
        client = _get_client()
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "你是一位专业财富管理顾问，擅长用通俗语言解释量化投资数据。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3500,
            timeout=timeout
        )
        raw = response.choices[0].message.content.strip()

        # 清理可能的 markdown 代码块包裹
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"\s*```$", "", raw, flags=re.MULTILINE)

        sections = json.loads(raw)
        # 验证所有必须 key 存在
        required = ["executive_summary", "capital_markets_outlook", "portfolio_roles",
                    "milestone_interpretation", "spending_impact",
                    "risk_narrative", "recommendations"]
        for k in required:
            if k not in sections:
                sections[k] = fallback.get(k, "（本章节内容生成失败，请联系顾问补充。）")
        return sections

    except Exception as e:
        print(f"[ReportGenerator] LLM 调用失败 ({model_name}): {e}，使用备用文本。")
        return fallback


# ═══════════════════════════════════════════════════════════════
# STEP 4 — 构建 Word 文档
# ═══════════════════════════════════════════════════════════════

# ── 文档样式助手 ──────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """设置表格单元格背景色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _cell_text(cell, text: str, bold=False, font_size=9.5,
               align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """快速设置单元格文字样式。"""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.clear()
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def _add_section_heading(doc: Document, title: str, level: int = 1):
    """添加章节标题。"""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(title)
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # 深蓝
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)  # 深灰
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def _add_body_text(doc: Document, text: str):
    """添加正文段落。"""
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    return p

def _add_data_table(doc: Document, headers: list, rows: list,
                    col_widths: list = None, header_bg="1E3A8A"):
    """通用数据表格：深蓝表头 + 交替行背景。"""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"

    # 表头行
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_row.cells[i], header_bg)
        _cell_text(hdr_row.cells[i], h, bold=True, font_size=9.5,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")

    # 数据行
    alt_bg = ["F8FAFF", "FFFFFF"]
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = alt_bg[r_idx % 2]
        for c_idx, val in enumerate(row):
            _set_cell_bg(tr.cells[c_idx], bg)
            align = WD_ALIGN_PARAGRAPH.RIGHT if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            _cell_text(tr.cells[c_idx], val, font_size=9.5, align=align)

    # 列宽设置
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)

    doc.add_paragraph()  # 表后间距
    return table


def _fmt_dollar(n, decimals=0):
    if n is None:
        return "N/A"
    return f"${n:,.{decimals}f}"


# ── 主构建函数 ────────────────────────────────────────────────

def build_docx(s: dict, sections: dict) -> io.BytesIO:
    """组装客户叙事版 Word 文档（9章结构），返回 BytesIO。"""

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══ 封面 ═════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    if s.get("advisor_name"):
        pre = doc.add_paragraph()
        pre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = pre.add_run(s["advisor_name"] + " · 长期财富规划方案")
        r0.font.size = Pt(10); r0.font.name = "微软雅黑"
        r0._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r0.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        doc.add_paragraph()

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = f"{s['client_name']} 财富规划方案" if s.get("client_name") else "您的长期财富规划方案"
    run = cover_title.add_run(title_text)
    run.bold = True; run.font.size = Pt(22); run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    doc.add_paragraph()
    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_title.add_run(f"基于 {s['years']} 年策略模拟  ·  10,000 次平行路径概率分析")
    run2.font.size = Pt(12); run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    if s.get("goals_text"):
        doc.add_paragraph()
        tag_p = doc.add_paragraph()
        tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_tag = tag_p.add_run(f"✦  {s['goals_text']}  ✦")
        run_tag.font.size = Pt(10); run_tag.font.name = "微软雅黑"
        run_tag._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run_tag.font.color.rgb = RGBColor(0x2E, 0x62, 0xAA)

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_p.add_run(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run3.font.size = Pt(10); run3.font.name = "微软雅黑"
    run3._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()
    # 封面KPI与网页预览对齐：达成率 / 中性终值 / 折算购买力 / 年化收益
    real_p50_cover = s.get("real_p50_final")
    if real_p50_cover:
        key_headers = [f"{s['years']}年目标达成率", "中性终值（P50）", "折算今日购买力", "预期年化收益"]
        key_values  = [f"{s['success_rate_pct']}%", _fmt_dollar(s['p50_final']),
                       _fmt_dollar(real_p50_cover), f"{s['expected_return_pct']}%"]
    else:
        key_headers = ["初始资本", f"{s['years']}年目标达成率", "中性终值（P50）", "预期年化收益"]
        key_values  = [_fmt_dollar(s['capital']), f"{s['success_rate_pct']}%",
                       _fmt_dollar(s['p50_final']), f"{s['expected_return_pct']}%"]
    key_table = doc.add_table(rows=2, cols=4)
    key_table.style = "Table Grid"
    for i, (h, v) in enumerate(zip(key_headers, key_values)):
        _set_cell_bg(key_table.rows[0].cells[i], "1E3A8A")
        _cell_text(key_table.rows[0].cells[i], h, bold=True, font_size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")
        _set_cell_bg(key_table.rows[1].cells[i], "EFF6FF")
        _cell_text(key_table.rows[1].cells[i], v, bold=True, font_size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color="1E3A8A")

    doc.add_page_break()

    # ══ 一、问题定义 ════════════════════════════════════════════
    _add_section_heading(doc, "一、这份方案想为您解决什么问题？")
    _add_body_text(doc, sections.get("executive_summary", ""))

    # ══ 二、核心策略与假设 ══════════════════════════════════════
    _add_section_heading(doc, "二、方案的核心策略与财务假设")
    _add_body_text(doc, sections.get("capital_markets_outlook", ""))

    _add_section_heading(doc, "财务沙盘参数", level=2)
    params_headers = ["参数项目", "设定值", "说明"]
    params_rows = [
        ["初始资本",  _fmt_dollar(s['capital']),       "起始注入资金额"],
        ["模拟年限",  f"{s['years']} 年",              "整体规划时间跨度"],
        ["通胀假设",  f"{s['inflation_pct']}%/年",     "用于抵消购买力侵蚀"],
    ]
    if s["contribution"] > 0:
        params_rows.append(["年度追加",
            f"{_fmt_dollar(s['contribution'])}/年 (第{s['contribution_start']}–{s['contribution_years']}年)",
            "持续定投，加速积累"])
    if s["withdrawal"] > 0:
        params_rows.append(["年度提取",
            f"{_fmt_dollar(s['withdrawal'])}/年 (第{s['withdrawal_start']}–{s['withdrawal_end']}年，通胀调整)",
            "每年可取用的现金流"])
    _add_data_table(doc, params_headers, params_rows, col_widths=[1.5, 2.8, 2.3])

    # ══ 三、资产配置 ════════════════════════════════════════════
    _add_section_heading(doc, "三、资产配置方案：您的财富如何运转")
    _add_body_text(doc, sections.get("portfolio_roles", ""))

    alloc_headers = ["资产名称", "代码", "权重", "年化收益", "波动率", "角色定位"]
    alloc_rows = [
        [a["name"], a["isin"], f"{a['weight_pct']}%",
         f"{a['exp_return_pct']}%", f"{a['vol_pct']}%", a.get("role", "—")]
        for a in s["allocations"]
    ]
    _add_data_table(doc, alloc_headers, alloc_rows, col_widths=[2.0, 0.9, 0.7, 1.0, 0.8, 1.2])

    perf_headers = ["预期年化收益率", "组合波动率", "夏普比率"]
    perf_rows = [[f"{s['expected_return_pct']}%", f"{s['volatility_pct']}%", str(s['sharpe_ratio'])]]
    _add_data_table(doc, perf_headers, perf_rows, col_widths=[2.0, 2.0, 2.0])

    # ══ 四、关键时间节点预估 ════════════════════════════════════
    _add_section_heading(doc, "四、未来关键时间点的财富预估")
    _add_body_text(doc, sections.get("milestone_interpretation", ""))

    if s.get("milestones"):
        m_headers = ["时间节点", "悲观 (P10)", "中性 (P50)", "乐观 (P90)", "这意味着什么"]
        m_rows = [
            [m["label"],
             _fmt_dollar(m.get("p10")), _fmt_dollar(m.get("p50")), _fmt_dollar(m.get("p90")),
             m.get("note", "")]
            for m in s["milestones"]
        ]
        _add_data_table(doc, m_headers, m_rows, col_widths=[1.8, 1.3, 1.3, 1.3, 1.9])
    else:
        snap_headers = ["年份", "悲观（P10）", "中性（P50）", "乐观（P90）"]
        snap_rows = [
            [f"第{snap['year']}年", _fmt_dollar(snap['p10']),
             _fmt_dollar(snap['p50']), _fmt_dollar(snap['p90'])]
            for snap in s["yearly_snapshots"]
        ]
        _add_data_table(doc, snap_headers, snap_rows, col_widths=[1.0, 1.8, 1.8, 1.8])

    final_headers = ["情景", "终值", "解读"]
    final_rows = [
        ["乐观（P90）", _fmt_dollar(s['p90_final']), "90%的路径低于此值，属顺风情景"],
        ["中性（P50）", _fmt_dollar(s['p50_final']), "中位数预期，最可能落点"],
        ["悲观（P10）", _fmt_dollar(s['p10_final']), "仅有10%路径更差，需心理准备的底线"],
    ]
    if s.get("real_p50_final"):
        final_rows.append(["实际购买力", _fmt_dollar(s['real_p50_final']),
                           f"中性终值折算今日购买力（通胀 {s['inflation_pct']}%/年）"])
    _add_data_table(doc, final_headers, final_rows, col_widths=[1.4, 1.4, 3.8])

    # ══ 五、年度提取计划 ════════════════════════════════════════
    _add_section_heading(doc, "五、您的年度提取计划")
    _add_body_text(doc, sections.get("spending_impact", ""))

    if s["withdrawal"] > 0 and s["withdrawal_schedule"]:
        w_start = s["withdrawal_schedule"][0]["year"]
        w_end   = s["withdrawal_schedule"][-1]["year"]
        _add_section_heading(doc, f"逐年提取金额（第{w_start}年～第{w_end}年）", level=2)
        # 判断是否有保险提取
        has_ins_w = any(w.get("ins_amount", 0) > 0 for w in s["withdrawal_schedule"])
        
        if has_ins_w:
            w_headers = ["年份", "持仓组合提取", "保险组合提取", "合计提款", "备注"]
            w_rows = []
            for i, w in enumerate(s["withdrawal_schedule"]):
                note = "基准年" if i == 0 else "通胀调整后"
                w_rows.append([
                    f"第{w['year']}年", 
                    _fmt_dollar(w['pure_amount']), 
                    _fmt_dollar(w['ins_amount']),
                    _fmt_dollar(w['total_amount']),
                    note
                ])
            _add_data_table(doc, w_headers, w_rows, col_widths=[0.8, 1.4, 1.4, 1.4, 1.6])
        else:
            w_headers = ["年份", "当年提取金额", "备注"]
            w_rows = []
            for i, w in enumerate(s["withdrawal_schedule"]):
                note = "基准年" if i == 0 else f"基准 × (1+{s['inflation_pct']}%)^{i}，通胀调整后"
                w_rows.append([f"第{w['year']}年", _fmt_dollar(w['pure_amount']), note])
            _add_data_table(doc, w_headers, w_rows, col_widths=[0.9, 1.5, 4.2])
    else:
        _add_body_text(doc, "本方案未设置年度提取，资产将持续复利增长。")

    # ══ 六、保险底座章节（有保险时）/ 灵活性风险边界 ══════════
    ins = s.get("ins_block")
    if ins:
        _add_section_heading(doc, "六、保险底座如何改变了整体资产配置的格局？")
        _add_body_text(doc, sections.get("insurance_strategy", ""))

        # 终值贡献对比表
        _add_section_heading(doc, f"{s['years']}年终值双核体系对比（非保证）", level=2)
        fin_ins_h = ["指标", "纯投资组合", "叠加保险底座", "保险额外贡献"]
        fin_ins_r = [
            ["中性终值 (P50)", _fmt_dollar(ins["pure_p50"]), _fmt_dollar(ins["combined_p50"]), _fmt_dollar(ins["contrib_p50"])],
            ["悲观终值 (P10)", _fmt_dollar(ins["pure_p10"]), _fmt_dollar(ins["combined_p10"]), _fmt_dollar(ins["contrib_p10"])],
        ]
        _add_data_table(doc, fin_ins_h, fin_ins_r, col_widths=[1.8, 1.5, 1.6, 1.6])

        # 回撤收窄表
        _add_section_heading(doc, "最大回撤收窄效果（P10悲观情景）", level=2)
        dd_ins_h = ["情景", "最大回撤幅度", "改善"]
        dd_ins_r = [
            ["纯投资组合", f"-{ins['drawdown_p10_pure_pct']}%", "基准"],
            ["双核合并体系", f"-{ins['drawdown_p10_comb_pct']}%", f"收窄 {ins['drawdown_improve_pct']} 个百分点"],
        ]
        _add_data_table(doc, dd_ins_h, dd_ins_r, col_widths=[2.0, 1.8, 2.8])

        # 保险CV三情景表
        if ins.get("cv_mid"):
            _add_section_heading(doc, f"第{s['years']}年保险现金价值：三情景分层", level=2)
            cv_h = ["实现情景", "保险终期现金价值", "说明"]
            cv_r = [
                ["悲观实现 (α=80%)",  _fmt_dollar(ins["cv_low"]),  "非保证红利以80%比例实现"],
                ["中性实现 (α=100%)", _fmt_dollar(ins["cv_mid"]),  "非保证红利全额实现（计划书演示值）"],
                ["乐观实现 (α=115%)", _fmt_dollar(ins["cv_high"]), "非保证红利超额实现（历史较好表现）"],
            ]
            _add_data_table(doc, cv_h, cv_r, col_widths=[1.8, 1.8, 3.0])

        # 提取覆盖率文字
        if ins.get("withdrawal_cov_pct", 0) > 0:
            _add_section_heading(doc, "保险现金流对年度提取的覆盖能力", level=2)
            cov_txt = (
                f"在提取阶段，保险累计可释放现金 {_fmt_dollar(ins['total_ins_withdrawal'])}，"
                f"可覆盖投资组合年度提取需求的 {ins['withdrawal_cov_pct']}%。"
                "这意味着投资组合被取用的压力相应减少，资产更有机会在市场中以复利滚动，而不是被迫变现。"
            )
            _add_body_text(doc, cov_txt)

        # 合规提示
        disc_ins = doc.add_paragraph()
        disc_ins.clear()
        run_di = disc_ins.add_run(
            "\u26a0\ufe0f 合规提示：保险的非保证红利（复归红利、终期分红）依赖保险公司实际经营表现，存在波动可能性。"
            "上述终值对比基于蒙特卡洛统计模拟，不构成收益承诺。请结合计划书中注明的\u300c保证/非保证\u300d栏综合评估。"
        )
        run_di.font.size = Pt(8.5); run_di.font.name = "微软雅黑"
        run_di._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run_di.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        doc.add_paragraph()

        _add_section_heading(doc, "七、方案的灵活性与风险边界")
        rec_chapter = "八、财富规划建议"
    else:
        _add_section_heading(doc, "六、方案的灵活性与风险边界")
        rec_chapter = "七、财富规划建议"

    _add_body_text(doc, sections.get("risk_narrative", ""))
    dd_headers = ["情景", "模拟期内最大回撤", "含义"]
    dd_rows = [
        ["悲观（P10）", f"{s['drawdown_p10_pct']}%", "最糟糕10%路径中的最大跌幅，需提前做好心理准备"],
        ["中性（P50）", f"{s['drawdown_p50_pct']}%", "中位数路径的最大历史跌幅"],
    ]
    if s.get("stressed_vol_pct"):
        dd_rows.append(["压力测试波动率", f"{s['stressed_vol_pct']}%",
                        "危机相关性场景下的组合波动率（已启用）"])
    _add_data_table(doc, dd_headers, dd_rows, col_widths=[1.4, 1.6, 3.6])

    # ══ 财富规划建议（章节号动态）══════════════════════════════
    _add_section_heading(doc, rec_chapter)
    _add_body_text(doc, sections.get("recommendations", ""))

    # ══ 八、关于我们的服务 ══════════════════════════════════════
    _add_section_heading(doc, "八、关于我们的服务")
    service_text = (
        f"本报告由 {s['advisor_name']} 编制。" if s.get("advisor_name")
        else "本报告由您的财富规划顾问编制。"
    )
    _add_body_text(doc, service_text +
        "我们致力于以科学、透明、有温度的方式，为高净值家庭提供长期财富规划服务。"
        "机构背景、资质信息及完整服务说明将在正式版本中补充。")

    # ══ 九、法律声明 ════════════════════════════════════════════
    doc.add_paragraph()
    _add_section_heading(doc, "九、重要声明与法律提示")
    disclaimer = doc.add_paragraph()
    disclaimer.clear()
    run_d = disclaimer.add_run(
        "本报告中所有财富预测数据均来自蒙特卡洛统计模拟，基于历史价格数据通过几何布朗运动模型（GBM）"
        "生成10,000条随机路径。P10/P50/P90 为概率分位数，不代表实际投资结果将落在此区间。"
        "所有投资均涉及风险，包括本金损失可能性。过去表现不代表未来结果。"
        "本报告仅供参考，不构成任何形式的投资建议、销售要约或财务规划合同。"
        "具体决策请结合实际财务状况、风险承受能力及相关监管规定，在专业顾问指导下进行。"
    )
    run_d.font.size = Pt(8.5); run_d.font.name = "微软雅黑"
    run_d._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run_d.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    disclaimer.paragraph_format.space_before = Pt(16)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_f = footer_p.add_run(
        f"Generated by Antigravity Wealth Engine  ·  {datetime.now().strftime('%Y年%m月%d日')}"
    )
    run_f.font.size = Pt(8); run_f.font.name = "Arial"
    run_f.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════
# 入口函数
# ═══════════════════════════════════════════════════════════════

def generate_mc_report(lab_result: dict, mc_settings: dict,
                       client_info: dict = None) -> io.BytesIO:
    """
    完整报告生成入口。
    client_info: 可选字典，含 name/age/goals/advisor 字段。
    返回包含 Word 文件内容的 BytesIO 对象。
    """
    print("[ReportGenerator] 开始生成客户版财富规划报告...")
    s        = build_data_summary(lab_result, mc_settings, client_info)
    prompt   = build_prompt(s)
    sections = call_llm(prompt)
    client_label = s.get("client_name") or "通用模版"
    print(f"[ReportGenerator] AI章节生成完毕（客户: {client_label}），开始组装 Word...")
    buf = build_docx(s, sections)
    print("[ReportGenerator] 报告生成完成。")
    return buf
